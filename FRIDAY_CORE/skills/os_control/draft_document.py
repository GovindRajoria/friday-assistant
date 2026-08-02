# skills/os_control/draft_document.py
from pathlib import Path

from core import llm_client
from docx import Document
from docx.shared import Pt


class DraftDocumentSkill:
    def __init__(self):
        self.manifest = {
            "name": "draft_document",
            "description": "Drafts, writes, or creates professional text documents (like letters, essays, or reports) and saves them as a Word file. 'topic' is the subject of the document.",
            "parameters": ["topic"]
        }

    def execute(self, params=None):
        if not params or "topic" not in params:
            return {"status": "error", "message": "I need a topic to draft the document."}
            
        topic = params["topic"]
        
        print(f"[*] FRIDAY is generating content for a document about '{topic}'...")
        
        try:
            # 1. Generate the content using Llama 3.1
            prompt = f"""
            Write a highly professional, well-formatted document regarding the following topic: {topic}.
            Do not include any conversational filler. Just output the content of the document.
            """
            
            # The shared client defaults to the low temperature the router wants;
            # prose written at 0.1 reads flat and repetitive, so this call opts
            # out of it explicitly.
            content = llm_client.chat([
                {'role': 'user', 'content': prompt}
            ], temperature=0.8).strip()
            
            # 2. Create the Word Document
            print("[*] Formatting and saving the document...")
            doc = Document()
            doc.add_heading(f'Research Draft: {topic}', 0)
            
            # Add the generated text
            paragraph = doc.add_paragraph(content)
            
            # Set a professional font size
            for run in paragraph.runs:
                run.font.size = Pt(11)
            
            filename = f"FRIDAY_Draft_{topic.replace(' ', '_')[:15]}.docx"
            
            # 3. Smart path logic: try Desktop, then Documents, then the cwd.
            # Path.home() resolves on every platform, unlike %USERPROFILE%.
            home = Path.home()
            potential_paths = [
                home / 'Desktop',
                home / 'Documents',
                Path.cwd(),
            ]

            for path in potential_paths:
                try:
                    if not path.exists():
                        continue

                    full_path = path / filename
                    doc.save(str(full_path))
                    return {"status": "success", "message": f"I have drafted the {topic} document and saved it to {full_path}."}
                except Exception:
                    continue
                    
            return {"status": "error", "message": "I couldn't find a writable directory to save the document."}

        except Exception as e:
            return {"status": "error", "message": f"I encountered an error while drafting the document: {str(e)}"}

def setup():
    return DraftDocumentSkill()
