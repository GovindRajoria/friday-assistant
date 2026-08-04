# skills/reading/read_document.py
"""Text out of a document on disk: PDF, DOCX, TXT, MD.

The gap this closes was embarrassing in practice — the assistant could fetch and
summarise a web page but could not open the file the operator was looking at, so
a resume PDF had to be pasted in by hand.

Two deliberate choices.

Every import is inside the method that needs it. `pypdf` is declared in
requirements.txt but CI does not install it, and `core/registry.py` drops a
skill whose module-scope import raises — silently, with one printed line. A
module-scope `import pypdf` would therefore make this skill *disappear* on any
machine missing the package rather than report a missing package, which is the
exact failure `skill_health` exists to surface.

Output is capped hard. A 200-page PDF is perhaps 400,000 characters and the
whole point is to put the text where a local model can reason about it, so a
document that does not fit is truncated *and says so*. A silent truncation would
have the model answering "the contract does not mention indemnity" about a
section it was never shown.
"""
from core.paths import allowed_roots, refusal, resolve_within

# Roughly 3,000 words. Enough for a resume, a README or a short paper; not
# enough to blow the prompt budget on one Observation.
MAX_CHARS = 12000
SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".json"}


class ReadDocumentSkill:
    def __init__(self):
        self.manifest = {
            "name": "read_document",
            "description": (
                "Reads the text out of a document file on this computer so you can "
                "summarise it, quote it, or answer questions from it. Handles PDF, DOCX, "
                "TXT and Markdown. Parameters: 'path', and optionally 'pages' for a PDF "
                "range like '1-5'. Use this for a FILE ON DISK — use read_webpage for a "
                "URL, and manage_files only to list a directory or delete something."
            ),
            "parameters": ["path", "pages"],
        }

    def execute(self, params=None):
        params = params or {}
        path_str = params.get("path")
        if not path_str:
            return {"status": "error", "message": "I need the path of a document to read."}

        roots = allowed_roots()
        resolved = resolve_within(str(path_str), roots)
        if resolved is None:
            return refusal(str(path_str))
        if not resolved.is_file():
            return {"status": "error", "message": f"'{resolved}' is not a readable file."}

        suffix = resolved.suffix.lower()
        if suffix not in SUPPORTED:
            return {
                "status": "error",
                "message": (f"I cannot read '{suffix or 'a file with no extension'}'. "
                            f"I handle: {', '.join(sorted(SUPPORTED))}."),
            }

        try:
            if suffix == ".pdf":
                text, note = self._read_pdf(resolved, params.get("pages"))
            elif suffix == ".docx":
                text, note = self._read_docx(resolved)
            else:
                text, note = self._read_plain(resolved)
        except ImportError as error:
            # Named plainly rather than swallowed: the fix is one pip install,
            # and a vague failure here sends someone debugging the document.
            return {
                "status": "error",
                "message": f"I am missing the library needed to read {suffix} files: {error}",
            }
        except Exception as error:                                    # noqa: BLE001
            return {"status": "error", "message": f"Could not read '{resolved.name}': {error}"}

        text = text.strip()
        if not text:
            return {
                "status": "success",
                "message": (f"'{resolved.name}' opened but contains no extractable text. "
                            "If it is a scanned document the pages are images — nothing to extract."),
                "data": {"path": str(resolved), "characters": 0},
            }

        truncated = len(text) > MAX_CHARS
        body = text[:MAX_CHARS]
        suffix_note = (f"\n\n[Truncated at {MAX_CHARS} characters of {len(text)}. "
                       "Ask for a specific section or page range for the rest.]" if truncated else "")
        return {
            "status": "success",
            "message": f"{resolved.name}{note}:\n\n{body}{suffix_note}",
            "data": {"path": str(resolved), "characters": len(text), "truncated": truncated},
        }

    def _read_pdf(self, path, pages):
        from pypdf import PdfReader

        reader = PdfReader(str(path))
        total = len(reader.pages)
        wanted = self._page_range(pages, total)
        chunks = []
        for index in wanted:
            try:
                chunks.append(reader.pages[index].extract_text() or "")
            except Exception as error:                                # noqa: BLE001
                # One malformed page should not lose the other 199.
                chunks.append(f"[page {index + 1} could not be extracted: {error}]")
        note = f" (pages {wanted[0] + 1}-{wanted[-1] + 1} of {total})" if wanted else f" ({total} pages)"
        return "\n\n".join(chunks), note

    @staticmethod
    def _page_range(pages, total) -> list[int]:
        """Parse '3' or '2-5' into zero-based indices, clamped to the document."""
        if not pages or not str(pages).strip():
            return list(range(total))
        text = str(pages).strip()
        try:
            if "-" in text:
                first, last = text.split("-", 1)
                start, end = int(first), int(last)
            else:
                start = end = int(text)
        except ValueError:
            return list(range(total))
        start = max(1, start)
        end = min(total, max(start, end))
        return list(range(start - 1, end))

    def _read_docx(self, path):
        import docx

        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        # Tables carry the content in a lot of real documents — a CV's dates, an
        # invoice's line items — and paragraph iteration alone silently skips them.
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts), ""

    def _read_plain(self, path):
        return path.read_text(encoding="utf-8", errors="replace"), ""


def setup():
    return ReadDocumentSkill()
